"""
Word Renderer — converts markdown text to a .docx file via python-docx.
"""
import os
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def run(input_data) -> dict:
    """
    Render markdown-like text to a .docx file.

    Args:
        input_data: dict with "content" (markdown text), "title", "output_path"
    """
    if isinstance(input_data, str):
        content = input_data
        title = "Document"
        output_path = "outputs/document.docx"
    elif isinstance(input_data, dict):
        content = input_data.get("content", input_data.get("text", ""))
        title = input_data.get("title", "Document")
        output_path = input_data.get("output_path", "outputs/document.docx")
    else:
        return {"status": "error", "result": "Invalid input"}

    content = str(content or "").strip()
    if not content:
        return {
            "status": "error",
            "result": "No content provided for Word rendering",
            "files": [],
        }

    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal_style = doc.styles["Normal"]
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.space_after = Pt(8)

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Parse markdown-like content
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            continue
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", line):
            text = re.sub(r"^\d+\.\s", "", line)
            doc.add_paragraph(text, style="List Number")
        else:
            doc.add_paragraph(line)

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)

    return {
        "status": "ok",
        "result": f"Word document saved: {output_path}",
        "files": [output_path],
    }
